# -*- coding: utf-8 -*-
"""Statement of Work documents."""
from collections import OrderedDict
from functools import partial
from itertools import chain
import re

from celery.exceptions import ImproperlyConfigured
from django.http import HttpResponse, HttpResponseNotFound
from django.shortcuts import get_object_or_404
from django.conf import settings
from docx import Document
from clients.models import ClientManifest
from projects.models import Project
from shared.templatetags.currency_filter import currency


import os.path

# This should maybe be a configurable pointer to an asset on S3 instead of
# being in the source directory, but the current version is smallish.
from shared.viewmodels import ProjectTargetSetViewModel

TEMPLATE_FILENAME = os.path.join(os.path.dirname(__file__), settings.SOW_TEMPLATE_DOCX)

DOCX_MIME_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'


def get_item(a_dict, key):
    return a_dict[key]


def generate_statement(project, template_file=None):
    if template_file is None:
        template_file = open(TEMPLATE_FILENAME, 'rb')

    doc = Document(template_file)

    table0 = doc.tables[0]
    table1 = doc.tables[1]
    table1.style = 'TableGrid'
    table2 = doc.tables[2]
    table3 = doc.tables[3]
    table4 = doc.tables[4]
    table5 = doc.tables[5]
    
    build_table(project, doc, table0, table1, table2, table3, table4, table5)
    replace_variables(project, doc)

    return doc


def build_table(project, doc, table0, table1, table2, table3, table4, table5):
    """Render the quote for this Project into this table.

    :type project: Project
    :type table: docx.table.Table
    """
    if not len(table1.columns) == 4:
        raise ImproperlyConfigured("SOW template's table is wrong size "
                                   "(%s rows x %s columns), expected 4 columns" %
                                   (len(table1.rows), len(table1.columns)))

    quote_summary_express = project.quote_express()
    quote_summary_standard = project.quote_standard()

    target_list = []

    table0_cells = table1.rows[0].cells
    table0_cells[0].text = ""
    table0_cells[0].paragraphs[0].add_run(u"Prepared expressly for %s" % (project.client.name,)).bold = True

    hdr_cells = table1.rows[1].cells
    hdr_cells[0].text = ""
    hdr_cells[0].paragraphs[0].add_run(u"%s" % (project.name,)).bold = True

    project_targets = ProjectTargetSetViewModel(project)

    # TODO: make sure rows/cells get styling

    for target in project_targets.targets:
        row = table1.add_row()
        row.style = "background-color: blue"
        row.cells[0].paragraphs[0].add_run(u"%s to %s" % (project.source_locale.description, target)).bold = True
        # row.cells[0].paragraphs[0].add_run(u"%s to %s" % (project.source_locale.description, target.description)).bold = True
        row.cells[1].paragraphs[0].add_run(u"Total").bold = True
        row.cells[2].paragraphs[0].add_run(u"MBD*").bold = True
        row.cells[3].paragraphs[0].add_run(u"Net Total").bold = True
        # target_list.append(target.description)
        
        for task in target.tasks:
            if not task.billable:
                continue
            row = table1.add_row()
            row.cells[0].paragraphs[0].text = task.service.service_type.description

            if task.project.is_standard_speed():
                price = task.net_price() if task.net_price() else 0
                raw_price = task.raw_price() if task.raw_price() else 0
                memory_bank_discount = task.mbd() if task.mbd() else 0
            else:
                price = task.express_net_price() if task.express_net_price() else 0
                raw_price = task.express_raw_price() if task.express_raw_price() else 0
                memory_bank_discount = task.express_mbd() if task.express_mbd() else 0

            row.cells[1].paragraphs[0].text = currency(price - ((raw_price * (memory_bank_discount*100))/100))
            row.cells[2].paragraphs[0].text = str('{:.1%}'.format(memory_bank_discount))
            row.cells[3].paragraphs[0].text = currency(price)

        row = table1.add_row()
        row.cells[0].paragraphs[0].add_run(u"Total").bold = True

        for target_quo in project.target_locales.all():
            target_details = get_item(project.target_price_details(), target.id)
            if target_quo.id == target.id:
                target_list.append(str(target_quo))
                if task.project.is_standard_speed():
                    row.cells[3].paragraphs[0].add_run(currency(target_details.target_price)).bold = True
                else:
                    row.cells[3].paragraphs[0].add_run(currency(target_details.target_express_price)).bold = True

    global num_targets
    num_targets = len(target_list)
    
    #table2 data
    hdr_cells = table2.rows[1].cells
    hdr_cells[1].text = currency(project.project_pricequote().price)
    hdr_cells[2].text = str('{0} business days'.format(int(abs(project.standard_duration()))))
    hdr_cells = table2.rows[2].cells
    hdr_cells[1].text = currency(project.project_pricequote().express_price)
    hdr_cells[2].text = str('{0} business days'.format(int(abs(project.express_duration()))))
    
    #Table3 data
    for asset in project.kit.source_files():
        row = table3.add_row()
        source_locales.append(asset.orig_name)
        row.cells[0].text = asset.orig_name

    #table4 data
    hdr_cells = table4.rows[0].cells
    hdr_cells[1].text = project.client.name
    hdr_cells = table4.rows[1].cells
    hdr_cells[1].text = project.name
    hdr_cells = table4.rows[2].cells
    hdr_cells[1].text = u'Same as source'
    hdr_cells = table4.rows[3].cells
    targets_list = ", ".join(target_list )
    hdr_cells[1].text = project.source_locale.description + u" > " + targets_list
    
    #table5 data
    hdr_cells = table5.rows[0].cells
    hdr_cells[1].text = currency(project.project_pricequote().price)
    hdr_cells[2].text = str('{0} business days'.format(int(abs(project.standard_duration()))))
    hdr_cells = table5.rows[1].cells
    hdr_cells[1].text = currency(project.project_pricequote().express_price)
    hdr_cells[2].text = str('{0} business days'.format(int(abs(project.express_duration()))))

variable_re = re.compile(r'(<NAME>|<CLIENT>|<JOB>|<EMAIL>|<PHONE>|<SOURCE_FILES>|<NUMBER_OF_TARGETS>|<VTP_URL>|<AE_NAME>|<AE_EMAIL>|<AE_PHONENUMBER>)')
bracket = lambda s: "[%s]" % (s,)
num_targets = 0
source_locales = []


def replace_variables(project, doc):
    """
    :type project: Project
    :type doc: Document
    """

    ctx = project_ctx(project)

    runs = chain.from_iterable(p.runs for p in doc.paragraphs)

    for run in runs:
        replace_in_run(run, ctx)


def project_ctx(project):
    sources_list = ", ".join(source_locales)
    del source_locales[:]
    return {
        '<NAME>': project.client_poc.get_full_name(),
        '<CLIENT>': project.client.name,
        '<JOB>': project.job_number,
        '<EMAIL>': project.client_poc.email,
        '<PHONE>': project.client.phone,
        '<SOURCE_FILES>': sources_list,
        '<NUMBER_OF_TARGETS>': str(num_targets),
        '<VTP_URL>': settings.BASE_URL,
        '<AE_NAME>': u" " if project.account_executive is None else u"%s %s," % (project.account_executive.first_name,project.account_executive.last_name),
        '<AE_EMAIL>': u" " if project.account_executive is None else u"%s," % project.account_executive.email,
        '<AE_PHONENUMBER>': u" " if project.account_executive is None else u"%s." % project.account_executive.phone
    }


def _sub_from_dict(ctx, match):
    text = match.group(1)
    new_text = ctx.get(text.strip())
    if new_text is None:
        # If we don't find it in the context, put the brackets back on so as
        # to return the original content.
        return bracket(text)
    return new_text


def replace_in_run(run, ctx):
    old_text = run.text
    new_text = variable_re.sub(partial(_sub_from_dict, ctx), old_text)   

    if old_text != new_text:
        run.text = new_text


def download_statement(request, pk):
    project = get_object_or_404(Project, id=pk)
    #SOW download option not available to client if is_sow_available is False. But is always available to VIA users.
    sow_download_option = True
    client_manifest = ClientManifest.objects.get(client__id=project.client.id)
    is_sow_available = client_manifest.is_sow_available and project.tasks_are_priced()
    if request.user.is_anonymous() or (request.user.is_client() and is_sow_available is False):
        sow_download_option = False
    if sow_download_option:
        doc = generate_statement(project)

        download_name = 'VTP_SOW-%s.docx' % (project.job_number,)

        response = HttpResponse(content_type=DOCX_MIME_TYPE)
        response['Content-Disposition'] = 'attachment; filename="%s"' % (
            download_name,)

        doc.save(response)

        return response
    return HttpResponseNotFound('<h1>No Page Here</h1>')